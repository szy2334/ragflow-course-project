"""Build deterministic evidence for the NeurIPS 2020 layout rules.

This script intentionally does not consume the repository's older fused JSON or
fusion builder.  PyMuPDF is used for native coordinates, fonts, baselines and
drawings; MinerU is used only for semantic roles and source text.  The output
contains rule-level observations so an LLM does not need to rediscover layout
geometry from a large raw extraction.

Example:
    python neurips_rule_evidence.py \
        --pdf NeurIPS-2020-bootstrap-your-own-latent-a-new-approach-to-self-supervised-learning-Paper.pdf \
        --mineru-json neur_minerU_result.json \
        --output neurips2020_byol.rule_evidence.json
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import statistics
import unicodedata
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Iterable


KNOWN_TYPES = {
    "title",
    "text",
    "ref_text",
    "image",
    "chart",
    "table",
    "image_caption",
    "chart_caption",
    "table_caption",
    "interline_equation",
    "equation",
    "code",
    "page_footnote",
    "header",
    "page_number",
}
SUBSET_PREFIX = re.compile(r"^[A-Z]{6}\+")
FIGURE_RE = re.compile(r"^(?:Figure|Fig\.)\s*(\d+)\s*[.:]", re.I)
TABLE_RE = re.compile(r"^Table\s*(\d+)\s*[.:]", re.I)
NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+){0,2})\.?\s+\D", re.I)
REFERENCE_NAMES = {"references", "bibliography"}
UNNUMBERED_H1_NAMES = {
    "broader impact",
    "broader impact statement",
    "acknowledgements",
    "acknowledgments",
    "acknowledgments and disclosure of funding",
    "references",
}
TIMES_COMPATIBLE = (
    "times",
    "nimbusromno9l",
    "liberationserif",
    "texgyretermes",
    "newtx",
    "stix",
    "xits",
)
SCHEMA_VERSION = "fused_layout_facts_v1"
EXPECTED_CAPTION_POSITION = {
    "figures": "caption_below",
    "tables": "caption_below",
}
RULE_DEFINITIONS = {
    "NIPS-01": "正文横向版心宽 396 pt；左边距 108 pt；正文区域起点距页顶约 72 pt。",
    "NIPS-02": "正文使用 Times New Roman 或 Times 兼容字体；主字号 10 pt；基线距离 11 pt。",
    "NIPS-03": "普通正文段落首行不缩进；相邻段落额外垂直间距约 5.5 pt。",
    "NIPS-04": "论文标题 17 pt、加粗、相对版心居中，并位于上下两条水平线之间。",
    "NIPS-05": "摘要左右各缩进约 36 pt；正文 10 pt、11 pt 行距、单段；Abstract 为 12 pt、加粗、居中。",
    "NIPS-06": "一级标题为 12 pt、加粗、相对正文区域左对齐。",
    "NIPS-07": "二级标题为 10 pt。",
    "NIPS-08": "三级标题为 10 pt。",
    "NIPS-09": "存在无编号的 References 一级标题，且在阅读顺序上位于正文之后。",
    "NIPS-10": "图、表均连续编号；图题位于图下；表题位于表下；图表相对局部或正文版心居中。",
}
RESULT_NAMES = {
    "compliant": "合规",
    "non_compliant": "不合规",
    "not_applicable": "不适用",
    "unverifiable": "无法可靠判断",
}


def normalize_text(value: Any) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = text.replace("\u00ad", "")
    text = re.sub(r"(?<=\w)-\s+(?=\w)", "", text)
    return re.sub(r"\s+", " ", text).strip()


def text_key(value: Any) -> str:
    return re.sub(r"[^\w]+", "", normalize_text(value).casefold(), flags=re.UNICODE)


def bbox(value: Any) -> list[float] | None:
    if not isinstance(value, (list, tuple)) or len(value) != 4:
        return None
    if not all(isinstance(item, (int, float)) and math.isfinite(item) for item in value):
        return None
    x0, y0, x1, y1 = map(float, value)
    return [x0, y0, x1, y1] if x1 > x0 and y1 > y0 else None


def bbox_union(values: Iterable[list[float] | None]) -> list[float] | None:
    boxes = [item for item in values if bbox(item)]
    if not boxes:
        return None
    return [
        min(item[0] for item in boxes),
        min(item[1] for item in boxes),
        max(item[2] for item in boxes),
        max(item[3] for item in boxes),
    ]


def intersection_ratio(inner: list[float] | None, outer: list[float] | None) -> float:
    if not bbox(inner) or not bbox(outer):
        return 0.0
    x0, y0 = max(inner[0], outer[0]), max(inner[1], outer[1])
    x1, y1 = min(inner[2], outer[2]), min(inner[3], outer[3])
    area = max(0.0, inner[2] - inner[0]) * max(0.0, inner[3] - inner[1])
    return max(0.0, x1 - x0) * max(0.0, y1 - y0) / area if area else 0.0


def horizontal_overlap(left: list[float], right: list[float]) -> float:
    overlap = max(0.0, min(left[2], right[2]) - max(left[0], right[0]))
    minimum = min(left[2] - left[0], right[2] - right[0])
    return overlap / minimum if minimum > 0 else 0.0


def similarity(left: str, right: str) -> float:
    a, b = text_key(left), text_key(right)
    if not a or not b:
        return 0.0
    if a in b or b in a:
        return min(len(a), len(b)) / max(len(a), len(b))
    # Avoid an additional dependency for a text match that only controls which
    # native lines are used for style measurement.
    common = sum(1 for char in set(a) if char in b)
    return common / max(len(set(a)), len(set(b)))


def weighted_median(values: list[tuple[float, int]]) -> float | None:
    if not values:
        return None
    ordered = sorted((float(value), max(1, int(weight))) for value, weight in values)
    half = sum(weight for _, weight in ordered) / 2
    total = 0
    for value, weight in ordered:
        total += weight
        if total >= half:
            return value
    return ordered[-1][0]


def style_for_lines(lines: list[dict[str, Any]], spans: dict[str, dict[str, Any]]) -> dict[str, Any]:
    selected = [spans[item] for line in lines for item in line.get("span_ids", []) if item in spans]
    sizes = [(float(item["font_size_pt"]), len(str(item.get("text") or ""))) for item in selected]
    total = sum(max(1, len(str(item.get("text") or ""))) for item in selected)
    bold = sum(max(1, len(str(item.get("text") or ""))) for item in selected if item.get("is_bold"))
    fonts = Counter(str(item.get("font_name") or "") for item in selected if item.get("font_name"))
    baselines = sorted(
        float(line["baseline_y"])
        for line in lines
        if isinstance(line.get("baseline_y"), (int, float))
    )
    gaps = [right - left for left, right in zip(baselines, baselines[1:]) if 5 <= right - left <= 30]
    return {
        "font_name": fonts.most_common(1)[0][0] if fonts else None,
        "font_size_pt": weighted_median(sizes),
        "bold_ratio": bold / total if total else None,
        "baseline_gap_pt": statistics.median(gaps) if gaps else None,
        "line_count": len(lines),
    }


def extract_pymupdf(pdf_path: Path) -> dict[str, Any]:
    import fitz

    document = fitz.open(pdf_path)
    pages: list[dict[str, Any]] = []
    spans: list[dict[str, Any]] = []
    lines: list[dict[str, Any]] = []
    text_blocks: list[dict[str, Any]] = []
    objects: list[dict[str, Any]] = []
    drawings: list[dict[str, Any]] = []
    try:
        for page_number, page in enumerate(document, start=1):
            pages.append(
                {
                    "page_number": page_number,
                    "width_pt": float(page.rect.width),
                    "height_pt": float(page.rect.height),
                    "rotation": int(page.rotation or 0),
                }
            )
            raw = page.get_text("dict")
            for block_index, raw_block in enumerate(raw.get("blocks", [])):
                if int(raw_block.get("type", 0)) != 0:
                    continue
                block_line_ids: list[str] = []
                block_text: list[str] = []
                for line_index, raw_line in enumerate(raw_block.get("lines", [])):
                    line_id = f"p{page_number}-l{block_index}-{line_index}"
                    line_span_ids: list[str] = []
                    parts: list[tuple[str, list[float]]] = []
                    baselines: list[float] = []
                    for span_index, raw_span in enumerate(raw_line.get("spans", [])):
                        value = str(raw_span.get("text") or "")
                        box = bbox(raw_span.get("bbox"))
                        if not value.strip() or not box:
                            continue
                        span_id = f"p{page_number}-s{block_index}-{line_index}-{span_index}"
                        raw_font = str(raw_span.get("font") or "") or None
                        flags = int(raw_span.get("flags") or 0)
                        origin = raw_span.get("origin")
                        baseline = float(origin[1]) if isinstance(origin, (list, tuple)) and len(origin) == 2 else None
                        if baseline is not None:
                            baselines.append(baseline)
                        spans.append(
                            {
                                "span_id": span_id,
                                "line_id": line_id,
                                "page_number": page_number,
                                "text": value,
                                "bbox": box,
                                "baseline_y": baseline,
                                "font_name": SUBSET_PREFIX.sub("", raw_font or "") or None,
                                "raw_font_name": raw_font,
                                "font_size_pt": float(raw_span.get("size") or 0.0),
                                "font_flags": flags,
                                "is_bold": "bold" in (raw_font or "").lower() or bool(flags & 16),
                                "is_italic": bool(flags & 2),
                            }
                        )
                        line_span_ids.append(span_id)
                        gap = box[0] - parts[-1][1][2] if parts else 0.0
                        if parts and gap >= 1 and not parts[-1][0].endswith(" ") and not value.startswith(" "):
                            parts.append((" ", [parts[-1][1][2], box[1], box[0], box[3]]))
                        parts.append((value, box))
                    if not line_span_ids:
                        continue
                    line_box = bbox(raw_line.get("bbox")) or bbox_union(item["bbox"] for item in spans if item["span_id"] in line_span_ids)
                    line_text = "".join(value for value, _ in parts).strip()
                    lines.append(
                        {
                            "line_id": line_id,
                            "native_block_id": f"p{page_number}-b{block_index}",
                            "page_number": page_number,
                            "text": line_text,
                            "bbox": line_box,
                            "baseline_y": statistics.median(baselines) if baselines else None,
                            "span_ids": line_span_ids,
                        }
                    )
                    block_line_ids.append(line_id)
                    block_text.append(line_text)
                if block_line_ids:
                    text_blocks.append(
                        {
                            "native_block_id": f"p{page_number}-b{block_index}",
                            "page_number": page_number,
                            "text": "\n".join(block_text),
                            "bbox": bbox(raw_block.get("bbox")),
                            "line_ids": block_line_ids,
                        }
                    )
            for image_index, image in enumerate(page.get_images(full=True), start=1):
                xref = int(image[0])
                for rect_index, rect in enumerate(page.get_image_rects(xref), start=1):
                    box = bbox((rect.x0, rect.y0, rect.x1, rect.y1))
                    if box:
                        objects.append({"object_id": f"p{page_number}-image-{xref}-{image_index}-{rect_index}", "object_type": "image", "page_number": page_number, "bbox": box})
            for drawing_index, drawing in enumerate(page.get_drawings(), start=1):
                record = {
                    "drawing_id": f"p{page_number}-drawing-{drawing_index}",
                    "page_number": page_number,
                    "bbox": bbox((drawing["rect"].x0, drawing["rect"].y0, drawing["rect"].x1, drawing["rect"].y1)) if drawing.get("rect") else None,
                    "width_pt": float(drawing.get("width") or 0.0),
                    "items": [],
                }
                for item_index, item in enumerate(drawing.get("items", []), start=1):
                    if not item:
                        continue
                    item_type = str(item[0])
                    child: dict[str, Any] = {"item_id": f"{record['drawing_id']}-{item_index}", "item_type": item_type}
                    if item_type == "l" and len(item) >= 3:
                        child.update({"start": [float(item[1].x), float(item[1].y)], "end": [float(item[2].x), float(item[2].y)], "width_pt": record["width_pt"]})
                    elif item_type == "re" and len(item) >= 2:
                        rect = item[1]
                        child["bbox"] = bbox((rect.x0, rect.y0, rect.x1, rect.y1))
                    record["items"].append(child)
                drawings.append(record)
                if record["bbox"]:
                    objects.append({"object_id": record["drawing_id"], "object_type": "vector_drawing", "page_number": page_number, "bbox": record["bbox"]})
    finally:
        document.close()
    return {"pages": pages, "spans": spans, "lines": lines, "text_blocks": text_blocks, "objects": objects, "drawings": drawings}


def node_text(node: dict[str, Any]) -> str:
    direct = node.get("text")
    if isinstance(direct, str) and direct.strip():
        return normalize_text(direct)
    parts: list[str] = []
    for line in node.get("lines", []) if isinstance(node.get("lines"), list) else []:
        for span in line.get("spans", []) if isinstance(line, dict) and isinstance(line.get("spans"), list) else []:
            if isinstance(span, dict) and isinstance(span.get("content"), str):
                parts.append(span["content"])
    return normalize_text(" ".join(parts))


def semantic_role(item_type: str, value: str) -> str:
    if FIGURE_RE.match(value):
        return "figure_caption"
    if TABLE_RE.match(value):
        return "table_caption"
    if item_type in {"image_caption", "chart_caption"}:
        return "figure_caption"
    if item_type == "table_caption":
        return "table_caption"
    if item_type == "title":
        return "title"
    if item_type == "ref_text":
        return "reference_entry"
    if item_type in {"image", "chart"}:
        return "figure_object"
    if item_type == "table":
        return "table_object"
    if item_type in {"equation", "interline_equation"}:
        return "display_formula"
    return "paragraph"


def parse_mineru(payload: Any, pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    source_sizes: dict[int, tuple[float, float]] = {}

    def emit(node: dict[str, Any], page: int, path: str, parent: str | None = None) -> None:
        item_type = str(node.get("type") or "unknown")
        if item_type not in KNOWN_TYPES:
            return
        raw_box = bbox(node.get("bbox"))
        target = next((item for item in pages if item["page_number"] == page), None)
        source = source_sizes.get(page)
        box = raw_box
        scale = [1.0, 1.0]
        if raw_box and source and target:
            scale = [target["width_pt"] / source[0], target["height_pt"] / source[1]]
            box = [raw_box[0] * scale[0], raw_box[1] * scale[1], raw_box[2] * scale[0], raw_box[3] * scale[1]]
        value = node_text(node)
        blocks.append(
            {
                "semantic_id": f"M{len(blocks) + 1:04d}",
                "parent_id": parent,
                "page_number": page,
                "mineru_type": item_type,
                "role": semantic_role(item_type, value),
                "heading_level": node.get("level", node.get("text_level")),
                "text": value,
                "bbox": box,
                "raw_bbox": raw_box,
                "coordinate_scale": scale,
            }
        )
        parent_id = blocks[-1]["semantic_id"]
        for index, child in enumerate(node.get("blocks", []) if isinstance(node.get("blocks"), list) else []):
            if isinstance(child, dict):
                emit(child, page, f"{path}.{index}", parent_id)

    if isinstance(payload, dict) and isinstance(payload.get("pdf_info"), list):
        for page_data in payload["pdf_info"]:
            page = int(page_data.get("page_idx", 0)) + 1
            size = page_data.get("page_size")
            if isinstance(size, (list, tuple)) and len(size) == 2:
                source_sizes[page] = (float(size[0]), float(size[1]))
        for page_data in payload["pdf_info"]:
            page = int(page_data.get("page_idx", 0)) + 1
            for index, node in enumerate(page_data.get("para_blocks") or page_data.get("preproc_blocks") or []):
                if isinstance(node, dict):
                    emit(node, page, str(index))
    elif isinstance(payload, list):
        for index, node in enumerate(payload):
            if isinstance(node, dict):
                page = int(node.get("page_idx", 0)) + 1
                emit(node, page, str(index))
    return blocks


def fuse_blocks(native: dict[str, Any], semantic: list[dict[str, Any]]) -> list[dict[str, Any]]:
    lines_by_page: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for line in native["lines"]:
        lines_by_page[line["page_number"]].append(line)
    spans = {item["span_id"]: item for item in native["spans"]}
    for item in semantic:
        page_lines = lines_by_page[item["page_number"]]
        box = item["bbox"]
        spatial = [line for line in page_lines if box and intersection_ratio(line["bbox"], [box[0] - 3, box[1] - 3, box[2] + 3, box[3] + 3]) >= 0.4]
        spatial.sort(key=lambda line: (line["bbox"][1], line["bbox"][0]))
        spatial_text = " ".join(line["text"] for line in spatial)
        native_blocks = [block for block in native["text_blocks"] if block["page_number"] == item["page_number"]]
        best_block = max(native_blocks, key=lambda block: similarity(item["text"], block["text"]), default=None)
        recovered = []
        if best_block and similarity(item["text"], best_block["text"]) >= 0.45:
            ids = set(best_block["line_ids"])
            recovered = [line for line in page_lines if line["line_id"] in ids]
        selected = spatial if spatial and similarity(item["text"], spatial_text) >= 0.25 else recovered
        item["native_line_ids"] = [line["line_id"] for line in selected]
        item["native_text"] = " ".join(line["text"] for line in selected).strip()
        item["native_bbox"] = bbox_union(line["bbox"] for line in selected)
        item["native_style"] = style_for_lines(selected, spans) if selected else {}
        item["text_similarity"] = similarity(item["text"], item["native_text"])
        item["match_status"] = "matched" if selected else "unmatched"
    return semantic


def mode(values: list[float], width: float = 2.0) -> float | None:
    if not values:
        return None
    buckets: dict[int, list[float]] = defaultdict(list)
    for value in values:
        buckets[round(value / width)].append(value)
    return statistics.median(max(buckets.values(), key=lambda group: (len(group), -statistics.pstdev(group))))


def body_lines(native: dict[str, Any], reference_page: int | None) -> list[dict[str, Any]]:
    spans = {item["span_id"]: item for item in native["spans"]}
    output = []
    for line in native["lines"]:
        if reference_page and line["page_number"] >= reference_page:
            continue
        if len(text_key(line["text"])) < 20 or line["bbox"][1] < 45:
            continue
        sizes = [spans[item]["font_size_pt"] for item in line["span_ids"] if item in spans]
        if not sizes or not 9.25 <= statistics.median(sizes) <= 10.75:
            continue
        output.append(line)
    return output


def reference_heading(blocks: list[dict[str, Any]]) -> dict[str, Any] | None:
    candidates = [item for item in blocks if item["role"] == "title" and text_key(item["text"]) in REFERENCE_NAMES]
    return min(candidates, key=lambda item: (item["page_number"], item["bbox"][1])) if candidates else None


def body_geometry(native: dict[str, Any], blocks: list[dict[str, Any]], ref: dict[str, Any] | None) -> dict[str, Any]:
    lines = body_lines(native, ref["page_number"] if ref else None)
    left = mode([line["bbox"][0] for line in lines])
    right = mode([line["bbox"][2] for line in lines])
    page_mins: list[float] = []
    for page in sorted({line["page_number"] for line in lines}):
        values = [line["bbox"][1] for line in lines if line["page_number"] == page and line["bbox"][1] < 100]
        if values:
            page_mins.append(min(values))
    top = mode(page_mins, 1.5)
    sizes: list[tuple[float, int]] = []
    fonts = Counter()
    spans = {item["span_id"]: item for item in native["spans"]}
    for line in lines:
        for span_id in line["span_ids"]:
            span = spans.get(span_id)
            if not span:
                continue
            weight = len(str(span.get("text") or ""))
            sizes.append((span["font_size_pt"], weight))
            fonts.update({span.get("font_name"): weight})
    baselines = defaultdict(list)
    for line in lines:
        if isinstance(line.get("baseline_y"), (int, float)):
            baselines[line["native_block_id"]].append(line["baseline_y"])
    gaps = [right_gap - left_gap for values in baselines.values() for left_gap, right_gap in zip(sorted(values), sorted(values)[1:]) if 5 <= right_gap - left_gap <= 30]
    return {
        "page_width_pt": native["pages"][0]["width_pt"] if native["pages"] else None,
        "page_height_pt": native["pages"][0]["height_pt"] if native["pages"] else None,
        "left_pt": left,
        "right_pt": right,
        "width_pt": right - left if left is not None and right is not None else None,
        "top_pt": top,
        "font_size_mode_pt": weighted_median(sizes),
        "font_name_mode": fonts.most_common(1)[0][0] if fonts else None,
        "font_name_times_compatible": bool(fonts and any(any(token in str(name).lower().replace("-", "") for token in TIMES_COMPATIBLE) for name in fonts)),
        "baseline_gap_median_pt": statistics.median(gaps) if gaps else None,
        "sample_line_count": len(lines),
        "sample_pages": sorted({line["page_number"] for line in lines}),
    }


def line_lookup(native: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {line["line_id"]: line for line in native["lines"]}


def merge_visual_lines(lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Merge PDF fragments that share a visual baseline.

    TeX frequently emits an inline bold heading and the following prose as two
    native line objects.  Treating the second fragment as the first line would
    manufacture a large indentation that is not visible in the PDF.
    """

    groups: list[list[dict[str, Any]]] = []
    for line in sorted(
        lines,
        key=lambda item: (
            float(item.get("baseline_y") or item["bbox"][3]),
            item["bbox"][0],
        ),
    ):
        baseline = float(line.get("baseline_y") or line["bbox"][3])
        if groups:
            previous = float(groups[-1][0].get("baseline_y") or groups[-1][0]["bbox"][3])
            if abs(baseline - previous) <= 1.0:
                groups[-1].append(line)
                continue
        groups.append([line])
    merged = []
    for group in groups:
        box = bbox_union(item["bbox"] for item in group)
        merged.append(
            {
                "bbox": box,
                "baseline_y": statistics.median(
                    float(item.get("baseline_y") or item["bbox"][3]) for item in group
                ),
                "text": " ".join(item["text"] for item in sorted(group, key=lambda item: item["bbox"][0])),
            }
        )
    return merged


def heading_depth(item: dict[str, Any]) -> int | None:
    match = NUMBERED_HEADING_RE.match(normalize_text(item["text"]))
    if match:
        return match.group(1).count(".") + 1
    if text_key(item["text"]) in {text_key(value) for value in UNNUMBERED_H1_NAMES}:
        return 1
    return None


def derive_front_matter(native: dict[str, Any], blocks: list[dict[str, Any]], geometry: dict[str, Any]) -> dict[str, Any]:
    title = next((item for item in blocks if item["role"] == "title" and item["page_number"] == 1), None)
    title_box = title["native_bbox"] or title["bbox"] if title else None
    center = (geometry["left_pt"] + geometry["right_pt"]) / 2 if geometry.get("left_pt") is not None else 306.0
    rules: list[dict[str, Any]] = []
    for drawing in native["drawings"]:
        if drawing["page_number"] != 1:
            continue
        for item in drawing["items"]:
            if item["item_type"] == "re" and item.get("bbox") and item["bbox"][2] - item["bbox"][0] >= 0.8 * (geometry.get("width_pt") or 396):
                rules.append({"kind": "filled_rectangle", "bbox": item["bbox"], "width_pt": item["bbox"][3] - item["bbox"][1]})
            elif item["item_type"] == "l" and item.get("start") and item.get("end") and abs(item["start"][1] - item["end"][1]) <= 0.5 and abs(item["end"][0] - item["start"][0]) >= 0.8 * (geometry.get("width_pt") or 396):
                rules.append({"kind": "line", "start": item["start"], "end": item["end"], "width_pt": item.get("width_pt")})
    ordered = sorted(rules, key=lambda item: item.get("bbox", [0, item.get("start", [0, 0])[1], 0, 0])[1] if item.get("bbox") else item["start"][1])
    def rule_y(item: dict[str, Any]) -> float:
        return float(item["bbox"][1] if item.get("bbox") else item["start"][1])

    between = [item for item in ordered if title_box and title_box[1] - 60 <= rule_y(item) < title_box[1]]
    below = [item for item in ordered if title_box and title_box[3] < rule_y(item) <= title_box[3] + 60]
    return {
        "title": title,
        "title_bbox": title_box,
        "title_center_x": (title_box[0] + title_box[2]) / 2 if title_box else None,
        "body_center_x": center,
        "title_alignment_delta_pt": abs((title_box[0] + title_box[2]) / 2 - center) if title_box else None,
        "horizontal_rules": ordered,
        "top_rule_candidates": between,
        "bottom_rule_candidates": below,
    }


def derive_abstract(blocks: list[dict[str, Any]], geometry: dict[str, Any], lines: dict[str, dict[str, Any]]) -> dict[str, Any] | None:
    heading = next((item for item in blocks if item["role"] == "title" and text_key(item["text"]) == "abstract"), None)
    if not heading or not heading["bbox"]:
        return None
    next_heading_y = min(
        (
            item["bbox"][1]
            for item in blocks
            if item["page_number"] == heading["page_number"]
            and item["role"] == "title"
            and item["bbox"]
            and item["bbox"][1] > heading["bbox"][3]
        ),
        default=float("inf"),
    )
    following = [item for item in blocks if item["page_number"] == heading["page_number"] and item["role"] == "paragraph" and item["bbox"] and heading["bbox"][3] <= item["bbox"][1] < next_heading_y]
    following.sort(key=lambda item: item["bbox"][1])
    if not following:
        return {"heading": heading, "body": None}
    line_items = [lines[line_id] for body in following for line_id in body["native_line_ids"] if line_id in lines]
    span_styles = [body["native_style"] for body in following if body.get("native_style")]
    body_box = bbox_union((body["native_bbox"] or body["bbox"]) for body in following)
    style = {
        "font_name": next((item.get("font_name") for item in span_styles if item.get("font_name")), None),
        "font_size_pt": statistics.median(item["font_size_pt"] for item in span_styles if item.get("font_size_pt") is not None) if any(item.get("font_size_pt") is not None for item in span_styles) else None,
        "bold_ratio": statistics.median(item["bold_ratio"] for item in span_styles if item.get("bold_ratio") is not None) if any(item.get("bold_ratio") is not None for item in span_styles) else None,
        "baseline_gap_pt": statistics.median(item["baseline_gap_pt"] for item in span_styles if item.get("baseline_gap_pt") is not None) if any(item.get("baseline_gap_pt") is not None for item in span_styles) else None,
    }
    left = geometry.get("left_pt")
    right = geometry.get("right_pt")
    return {
        "heading": heading,
        "heading_style": heading.get("native_style") or {},
        "heading_alignment_delta_pt": abs(
            ((heading.get("native_bbox") or heading["bbox"])[0] + (heading.get("native_bbox") or heading["bbox"])[2]) / 2
            - (left + right) / 2
        ) if left is not None and right is not None else None,
        "body": following,
        "body_bbox": body_box,
        "paragraph_count": len(following),
        "left_indent_pt": body_box[0] - left if body_box and left is not None else None,
        "right_indent_pt": right - body_box[2] if body_box and right is not None else None,
        "style": style,
        "line_count": len(line_items),
    }


def paragraph_metrics(blocks: list[dict[str, Any]], geometry: dict[str, Any], lines: dict[str, dict[str, Any]]) -> dict[str, Any]:
    candidates = [item for item in blocks if item["role"] == "paragraph" and item["native_bbox"] and 1 <= item["page_number"]]
    candidates = [item for item in candidates if item["native_style"].get("font_size_pt") and abs(item["native_style"]["font_size_pt"] - 10) <= 0.6 and item["native_bbox"][0] <= (geometry.get("left_pt") or 108) + 8]
    first_deltas: list[float] = []
    by_page: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for item in candidates:
        selected = merge_visual_lines([lines[line_id] for line_id in item["native_line_ids"] if line_id in lines])
        selected = [line for line in selected if line["bbox"][0] <= (geometry.get("left_pt") or 108) + 4 and len(text_key(line["text"])) >= 8]
        if len(selected) >= 2:
            first_deltas.append(selected[0]["bbox"][0] - statistics.median(line["bbox"][0] for line in selected[1:]))
        by_page[item["page_number"]].append(item)
    extra_gaps: list[float] = []
    for page_items in by_page.values():
        page_items.sort(key=lambda item: item["native_bbox"][1])
        for left_item, right_item in zip(page_items, page_items[1:]):
            if right_item["native_bbox"][1] < left_item["native_bbox"][3]:
                continue
            left_lines = merge_visual_lines([lines[line_id] for line_id in left_item["native_line_ids"] if line_id in lines])
            right_lines = merge_visual_lines([lines[line_id] for line_id in right_item["native_line_ids"] if line_id in lines])
            left_baselines = [line["baseline_y"] for line in left_lines if isinstance(line.get("baseline_y"), (int, float))]
            right_baselines = [line["baseline_y"] for line in right_lines if isinstance(line.get("baseline_y"), (int, float))]
            if left_baselines and right_baselines:
                gap = min(right_baselines) - max(left_baselines)
                if 13 <= gap <= 22:
                    extra_gaps.append(gap - 11.0)
    return {
        "paragraph_count_sampled": len(candidates),
        "first_line_indent_median_pt": statistics.median(first_deltas) if first_deltas else None,
        "first_line_indent_abs_p90_pt": sorted(abs(value) for value in first_deltas)[max(0, int(len(first_deltas) * 0.9) - 1)] if first_deltas else None,
        "paragraph_extra_gap_median_pt": statistics.median(extra_gaps) if extra_gaps else None,
        "paragraph_extra_gap_samples": len(extra_gaps),
    }


def caption_records(blocks: list[dict[str, Any]], geometry: dict[str, Any]) -> dict[str, Any]:
    captions = {"figures": [], "tables": []}
    objects = [item for item in blocks if item["role"] in {"figure_object", "table_object"} and item["bbox"]]
    for item in blocks:
        match = FIGURE_RE.match(item["text"])
        kind = "figures"
        if not match:
            match = TABLE_RE.match(item["text"])
            kind = "tables"
        if not match or not item["bbox"]:
            continue
        number = int(match.group(1))
        object_role = "figure_object" if kind == "figures" else "table_object"
        candidates: list[tuple[float, dict[str, Any], str]] = []
        for obj in objects:
            if obj["role"] != object_role or obj["page_number"] != item["page_number"]:
                continue
            overlap = horizontal_overlap(item["bbox"], obj["bbox"])
            if overlap < 0.1:
                continue
            if obj["bbox"][3] <= item["bbox"][1] + 4:
                direction, gap = "caption_below", item["bbox"][1] - obj["bbox"][3]
            elif item["bbox"][3] <= obj["bbox"][1] + 4:
                direction, gap = "caption_above", obj["bbox"][1] - item["bbox"][3]
            else:
                continue
            if gap <= 180:
                candidates.append((gap - overlap * 10, obj, direction))
        if not candidates:
            captions[kind].append(
                {
                    "number": number,
                    "caption": item,
                    "object_ids": [],
                    "position": None,
                    "center_delta_pt": None,
                    "body_center_delta_pt": None,
                    "alignment_scope": None,
                    "centered": False,
                }
            )
            continue
        best_score, _, direction = min(candidates, key=lambda value: value[0])
        selected = [obj for _, obj, obj_direction in candidates if obj_direction == direction and abs(_ - best_score) < 80]
        selected = selected or [min(candidates, key=lambda value: value[0])[1]]
        object_box = bbox_union(obj["bbox"] for obj in selected)
        cap_center = (item["bbox"][0] + item["bbox"][2]) / 2
        object_center = (object_box[0] + object_box[2]) / 2 if object_box else None
        body_center = (geometry["left_pt"] + geometry["right_pt"]) / 2
        body_width = geometry["right_pt"] - geometry["left_pt"]
        center_delta = abs(cap_center - object_center) if object_center is not None else None
        body_center_delta = abs(object_center - body_center) if object_center is not None else None
        is_cross_width = bool(object_box and object_box[2] - object_box[0] >= body_width * 0.55)
        captions[kind].append({
            "number": number,
            "caption": item,
            "object_ids": [obj["semantic_id"] for obj in selected],
            "object_bbox": object_box,
            "position": direction,
            "gap_pt": min(candidates, key=lambda value: value[0])[1]["bbox"][1] - item["bbox"][3] if direction == "caption_above" else item["bbox"][1] - min(candidates, key=lambda value: value[0])[1]["bbox"][3],
            "center_delta_pt": center_delta,
            "body_center_delta_pt": body_center_delta,
            "alignment_scope": "body_width" if is_cross_width else "local_object",
            "centered": bool(center_delta is not None and center_delta <= 10 and (not is_cross_width or body_center_delta <= 10)),
        })
    for kind, records in captions.items():
        records.sort(key=lambda item: (item["caption"]["page_number"], item["caption"]["bbox"][1]))
        numbers = [item["number"] for item in records]
        captions[kind] = {
            "count": len(records),
            "numbers": numbers,
            "unique": len(numbers) == len(set(numbers)),
            "continuous": numbers == list(range(1, len(numbers) + 1)) and len(numbers) == len(set(numbers)) if numbers else None,
            "items": records,
        }
    return captions


def result(rule_id: str, title: str, status: str, evidence: dict[str, Any], confidence: float = 0.9) -> dict[str, Any]:
    return {"rule_id": rule_id, "title": title, "result": status, "confidence": round(confidence, 3), "evidence": evidence}


def measured_status(available: bool, passed: bool) -> str:
    if not available:
        return "unverifiable"
    return "compliant" if passed else "non_compliant"


def round_floats(value: Any) -> Any:
    if isinstance(value, float):
        return round(value, 4)
    if isinstance(value, dict):
        return {key: round_floats(item) for key, item in value.items()}
    if isinstance(value, list):
        return [round_floats(item) for item in value]
    return value


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def serialize_fused_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    for index, item in enumerate(blocks, start=1):
        style = item.get("native_style") or {}
        is_object = item["role"] in {"figure_object", "table_object"}
        native_box = item.get("native_bbox")
        selected_box = item.get("bbox") if is_object else native_box or item.get("bbox")
        selected_source = (
            "mineru_semantic_container_cross_checked_by_pymupdf"
            if is_object
            else "pymupdf_text_geometry"
            if native_box
            else "mineru"
        )
        confidence = float(item.get("text_similarity") or 0.0)
        if item.get("match_status") == "matched" and confidence < 0.65:
            confidence = 0.75 if is_object else 0.65
        output.append(
            {
                "fused_block_id": f"fused-{index:05d}",
                "page_number": item["page_number"],
                "semantic_role": item["role"],
                "heading_level": heading_depth(item) or item.get("heading_level"),
                "text": item["text"],
                "bbox": selected_box,
                "style": {
                    "dominant_font": style.get("font_name"),
                    "font_size_median_pt": style.get("font_size_pt"),
                    "bold_character_ratio": style.get("bold_ratio"),
                    "baseline_gap_median_pt": style.get("baseline_gap_pt"),
                },
                "match": {
                    "confidence": confidence,
                    "text_similarity": item.get("text_similarity"),
                    "location_method": (
                        "mineru_container_cross_checked_by_pymupdf"
                        if is_object
                        else "pymupdf_text_match"
                        if native_box
                        else "mineru_only"
                    ),
                    "status": item.get("match_status"),
                },
                "bbox_validation": {
                    "geometry_valid": bool(item.get("bbox")),
                    "selected_bbox_source": selected_source,
                    "mineru_pymupdf_overlap": intersection_ratio(item.get("bbox"), native_box),
                },
                "source_ids": {
                    "mineru_block_ids": [item["semantic_id"]],
                    "pymupdf_line_ids": item.get("native_line_ids", []),
                },
                "field_provenance": {
                    "text": "mineru",
                    "semantic_role": "mineru",
                    "bbox": selected_source,
                    "style": "pymupdf" if style else None,
                },
                "alternatives": {
                    "mineru_type": item.get("mineru_type"),
                    "mineru_heading_level": item.get("heading_level"),
                    "mineru_text": item.get("text"),
                    "mineru_bbox": item.get("bbox"),
                    "pymupdf_text": item.get("native_text"),
                    "pymupdf_text_bbox": native_box,
                },
            }
        )
    return output


def build_derived_facts(
    native: dict[str, Any], report: dict[str, Any], blocks: list[dict[str, Any]]
) -> dict[str, Any]:
    geometry = report["body_geometry"]
    reference = report["references"].get("heading")
    page_sizes = {
        (round(page["width_pt"], 2), round(page["height_pt"], 2))
        for page in native["pages"]
    }
    return {
        "page_geometry": {
            "page_count": len(native["pages"]),
            "pages": native["pages"],
            "all_pages_same_size": len(page_sizes) == 1,
        },
        "document_structure": {
            "page_count": len(native["pages"]),
            "references_start": (
                {
                    "page_number": reference["page_number"],
                    "bbox": reference["native_bbox"] or reference["bbox"],
                    "evidence_id": reference["semantic_id"],
                }
                if reference
                else None
            ),
            "reference_entry_count": sum(item["role"] == "reference_entry" for item in blocks),
        },
        "column_geometry": {
            "column_count": 1,
            "text_left_pt": geometry.get("left_pt"),
            "text_right_pt": geometry.get("right_pt"),
            "text_width_pt": geometry.get("width_pt"),
            "body_font_size_mode_pt": geometry.get("font_size_mode_pt"),
            "sample_line_count": geometry.get("sample_line_count"),
            "sample_pages": geometry.get("sample_pages"),
            "derivation": "repeated native text-line boundary modes",
            "confidence": 0.95,
        },
        "body_geometry": report["body_geometry"],
        "front_matter": report["front_matter"],
        "abstract": report["abstract"],
        "paragraph_metrics": report["paragraph_metrics"],
        "headings": report["headings"],
        "captions": report["captions"],
        "references": report["references"],
        "rule_assessments": report["rules"],
    }


def build_rule_evidence_index(rule_assessments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    paths = {
        "NIPS-01": ["derived_facts.body_geometry", "derived_facts.page_geometry"],
        "NIPS-02": ["derived_facts.body_geometry", "native_facts.spans", "native_facts.lines"],
        "NIPS-03": ["derived_facts.paragraph_metrics"],
        "NIPS-04": ["derived_facts.front_matter", "native_facts.drawings"],
        "NIPS-05": ["derived_facts.abstract"],
        "NIPS-06": ["derived_facts.headings"],
        "NIPS-07": ["derived_facts.headings"],
        "NIPS-08": ["derived_facts.headings"],
        "NIPS-09": ["derived_facts.references", "derived_facts.document_structure"],
        "NIPS-10": ["derived_facts.captions"],
    }
    return [
        {
            "rule_key": item["rule_id"],
            "evidence_paths": paths[item["rule_id"]],
            "assessment": item["result"],
        }
        for item in rule_assessments
    ]


def compact_number(value: Any) -> str:
    if not isinstance(value, (int, float)):
        return "缺失"
    return f"{float(value):.2f}"


def rule_observation(rule_id: str, derived: dict[str, Any]) -> str:
    geometry = derived["body_geometry"]
    if rule_id == "NIPS-01":
        return (
            f"左={compact_number(geometry.get('left_pt'))} pt，"
            f"右={compact_number(geometry.get('right_pt'))} pt，"
            f"宽={compact_number(geometry.get('width_pt'))} pt，"
            f"顶部={compact_number(geometry.get('top_pt'))} pt"
        )
    if rule_id == "NIPS-02":
        return (
            f"字体={geometry.get('font_name_mode') or '缺失'}，"
            f"字号={compact_number(geometry.get('font_size_mode_pt'))} pt，"
            f"基线距={compact_number(geometry.get('baseline_gap_median_pt'))} pt"
        )
    if rule_id == "NIPS-03":
        value = derived["paragraph_metrics"]
        return (
            f"首行缩进中位数={compact_number(value.get('first_line_indent_median_pt'))} pt，"
            f"段间额外距离={compact_number(value.get('paragraph_extra_gap_median_pt'))} pt"
        )
    if rule_id == "NIPS-04":
        value = derived["front_matter"]
        style = (value.get("title") or {}).get("native_style") or {}
        return (
            f"字号={compact_number(style.get('font_size_pt'))} pt，"
            f"粗体比例={compact_number(style.get('bold_ratio'))}，"
            f"中心偏差={compact_number(value.get('title_alignment_delta_pt'))} pt，"
            f"上下横线={len(value.get('top_rule_candidates') or [])}/"
            f"{len(value.get('bottom_rule_candidates') or [])}"
        )
    if rule_id == "NIPS-05":
        value = derived.get("abstract") or {}
        style = value.get("style") or {}
        heading_style = value.get("heading_style") or {}
        return (
            f"左右缩进={compact_number(value.get('left_indent_pt'))}/"
            f"{compact_number(value.get('right_indent_pt'))} pt，"
            f"正文={compact_number(style.get('font_size_pt'))}/"
            f"{compact_number(style.get('baseline_gap_pt'))} pt，"
            f"Abstract={compact_number(heading_style.get('font_size_pt'))} pt，"
            f"段落数={value.get('paragraph_count', '缺失')}"
        )
    if rule_id in {"NIPS-06", "NIPS-07", "NIPS-08"}:
        depth = int(rule_id[-1]) - 5
        values = [item for item in derived["headings"] if item["depth"] == depth]
        sizes = [
            item["style"].get("font_size_pt")
            for item in values
            if item["style"].get("font_size_pt") is not None
        ]
        return (
            f"检出={len(values)}，字号中位数="
            f"{compact_number(statistics.median(sizes) if sizes else None)} pt"
        )
    if rule_id == "NIPS-09":
        count = derived["document_structure"].get("reference_entry_count")
        heading = derived["references"].get("heading") or {}
        return f"标题={heading.get('text') or '缺失'}，参考文献条目={count}"
    captions = derived["captions"]
    figure_positions = [item.get("position") for item in captions["figures"]["items"]]
    table_positions = [item.get("position") for item in captions["tables"]["items"]]
    return (
        f"图编号={captions['figures']['numbers']}，图题位置={figure_positions}；"
        f"表编号={captions['tables']['numbers']}，表题位置={table_positions}"
    )


def write_rule_document(path: Path, output: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(10)

    def format_run(run: Any, size: float, *, bold: bool = False, color: str | None = None) -> None:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def add_section(title_text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_with_next = True
        format_run(paragraph.add_run(title_text), 13, bold=True, color="111827")

    def add_bullet(text: str) -> None:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        format_run(paragraph.add_run(text), 10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    format_run(title.add_run("NeurIPS 2020 可验证格式规则"), 18, bold=True, color="1F2937")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    format_run(
        subtitle.add_run("基于 MinerU 与 PyMuPDF 融合事实 JSON 的确定性检查版本"),
        10,
        color="4B5563",
    )

    add_section("1. 正文区域几何（NIPS-01）")
    add_bullet("正文横向版心宽度：396 ± 3 pt。")
    add_bullet("正文左边界：108 ± 3 pt。")
    add_bullet("正文区域起点距页顶：72 ± 3 pt。")

    add_section("2. 正文字体与行距（NIPS-02）")
    add_bullet("正文主字体必须为 Times New Roman 或属于以下 Times 兼容字体白名单。")
    add_bullet("正文主字号：9.75–10.25 pt。")
    add_bullet("正文相邻行基线距离中位数：10.5–11.5 pt。")
    whitelist_heading = document.add_paragraph()
    whitelist_heading.paragraph_format.space_before = Pt(5)
    whitelist_heading.paragraph_format.space_after = Pt(4)
    format_run(whitelist_heading.add_run("Times 兼容字体白名单"), 10.5, bold=True)
    whitelist = document.add_table(rows=1, cols=1)
    whitelist.style = "Normal Table"
    whitelist.autofit = True
    cell_paragraph = whitelist.cell(0, 0).paragraphs[0]
    cell_paragraph.paragraph_format.space_after = Pt(0)
    whitelist_names = (
        "Times*\nNimbusRomNo9L*\nLiberationSerif*\nTeXGyreTermes*\n"
        "NewTX*\nSTIX*\nXITS*"
    )
    whitelist_run = cell_paragraph.add_run(whitelist_names)
    whitelist_run.font.name = "Consolas"
    whitelist_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Consolas")
    whitelist_run.font.size = Pt(9)

    add_section("3. 段落格式（NIPS-03）")
    add_bullet("普通正文段落首行缩进中位数的绝对值不超过 1.5 pt。")
    add_bullet("相邻段落额外垂直间距中位数：3.5–7.5 pt。")

    add_section("4. 论文标题（NIPS-04）")
    add_bullet("标题字号：16.5–17.5 pt。")
    add_bullet("标题粗体字符比例至少为 0.9。")
    add_bullet("标题中心与正文区域中心偏差不超过 8 pt。")
    add_bullet("标题上方和下方必须各检测到至少一条水平线。")

    add_section("5. 摘要（NIPS-05）")
    add_bullet("摘要正文相对正文区域左右各额外缩进：31–41 pt。")
    add_bullet("摘要正文字号：9.5–10.5 pt。")
    add_bullet("摘要正文相邻行基线距离：10.25–11.75 pt。")
    add_bullet("摘要正文由一个段落组成。")
    add_bullet("Abstract 标题字号：11.5–12.5 pt。")
    add_bullet("Abstract 标题粗体字符比例至少为 0.9。")
    add_bullet("Abstract 标题中心与正文区域中心偏差不超过 8 pt。")

    add_section("6. 一级标题（NIPS-06）")
    add_bullet("一级标题字号：11.5–12.5 pt。")
    add_bullet("一级标题粗体字符比例至少为 0.8。")
    add_bullet("一级标题左边界与正文区域左边界偏差不超过 3 pt。")

    add_section("7. 二级标题（NIPS-07）")
    add_bullet("二级标题字号：9.5–10.5 pt。")
    add_bullet("文档中不存在二级标题时，本规则输出不适用。")

    add_section("8. 三级标题（NIPS-08）")
    add_bullet("三级标题字号：9.5–10.5 pt。")
    add_bullet("文档中不存在三级标题时，本规则输出不适用。")

    add_section("9. 参考文献标题（NIPS-09）")
    add_bullet("必须存在无编号的 References 一级标题。")
    add_bullet("References 标题在文档阅读顺序上位于正文之后。")
    add_bullet("References 标题之后必须存在可识别的参考文献条目。")

    add_section("10. 图表结构（NIPS-10）")
    add_bullet("图号和表号分别从 1 开始连续编号，且编号不得重复。")
    add_bullet("图题必须位于图形下方。")
    add_bullet("表题必须位于表格下方。")
    add_bullet("图题和表题相对所属栏或跨栏正文区域居中，中心偏差不超过 10 pt。")
    add_bullet("文档中既不存在图也不存在表时，本规则输出不适用。")

    document.save(path)


def run_self_tests() -> None:
    assert heading_depth({"text": "3 Headings: first level"}) == 1
    assert heading_depth({"text": "3.1 Headings: second level"}) == 2
    assert heading_depth({"text": "3.1.1 Headings: third level"}) == 3
    assert heading_depth({"text": "References"}) == 1
    merged = merge_visual_lines(
        [
            {"bbox": [108.0, 100.0, 220.0, 112.0], "baseline_y": 110.0, "text": "Bold heading"},
            {"bbox": [225.0, 100.0, 504.0, 112.0], "baseline_y": 110.0, "text": "body text"},
            {"bbox": [108.0, 111.0, 504.0, 123.0], "baseline_y": 121.0, "text": "next line"},
        ]
    )
    assert len(merged) == 2
    assert merged[0]["bbox"] == [108.0, 100.0, 504.0, 112.0]
    duplicate_blocks = [
        {"semantic_id": "M1", "role": "figure_caption", "page_number": 1, "text": "Figure 1: first", "bbox": [108.0, 200.0, 300.0, 212.0]},
        {"semantic_id": "M2", "role": "figure_caption", "page_number": 2, "text": "Figure 1: duplicate", "bbox": [108.0, 200.0, 300.0, 212.0]},
    ]
    facts = caption_records(duplicate_blocks, {"left_pt": 108.0, "right_pt": 504.0})
    assert facts["figures"]["unique"] is False
    assert facts["figures"]["continuous"] is False
    assert measured_status(False, False) == "unverifiable"
    assert measured_status(True, False) == "non_compliant"
    assert measured_status(True, True) == "compliant"
    assert EXPECTED_CAPTION_POSITION == {
        "figures": "caption_below",
        "tables": "caption_below",
    }


def evaluate(native: dict[str, Any], blocks: list[dict[str, Any]]) -> dict[str, Any]:
    ref = reference_heading(blocks)
    geometry = body_geometry(native, blocks, ref)
    lines = line_lookup(native)
    front = derive_front_matter(native, blocks, geometry)
    abstract = derive_abstract(blocks, geometry, lines)
    paragraphs = paragraph_metrics(blocks, geometry, lines)
    headings: list[dict[str, Any]] = []
    for item in blocks:
        if item["role"] != "title" or item["page_number"] != 1 and item.get("text") == "":
            continue
        depth = heading_depth(item)
        if depth:
            headings.append({"semantic_id": item["semantic_id"], "text": item["text"], "depth": depth, "page_number": item["page_number"], "style": item["native_style"], "bbox": item["native_bbox"] or item["bbox"]})
    captions = caption_records(blocks, geometry)
    body_font_ok = geometry["font_size_mode_pt"] is not None and abs(geometry["font_size_mode_pt"] - 10) <= 0.25
    baseline_ok = geometry["baseline_gap_median_pt"] is not None and abs(geometry["baseline_gap_median_pt"] - 11) <= 0.5
    font_ok = bool(geometry["font_name_times_compatible"])
    title = front.get("title")
    title_style = title.get("native_style", {}) if title else {}
    title_box = front.get("title_bbox")
    title_rules_ok = len(front["top_rule_candidates"]) >= 1 and len(front["bottom_rule_candidates"]) >= 1
    h1 = [item for item in headings if item["depth"] == 1]
    h2 = [item for item in headings if item["depth"] == 2]
    h3 = [item for item in headings if item["depth"] == 3]
    geometry_available = all(geometry.get(key) is not None for key in ("left_pt", "right_pt", "width_pt", "top_pt"))
    geometry_ok = geometry_available and abs(geometry["width_pt"] - 396) <= 3 and abs(geometry["left_pt"] - 108) <= 3 and abs(geometry["top_pt"] - 72) <= 3
    body_style_available = geometry.get("font_size_mode_pt") is not None and geometry.get("baseline_gap_median_pt") is not None and geometry.get("font_name_mode") is not None
    paragraph_available = paragraphs["first_line_indent_median_pt"] is not None and paragraphs["paragraph_extra_gap_median_pt"] is not None
    paragraph_ok = paragraph_available and abs(paragraphs["first_line_indent_median_pt"]) <= 1.5 and abs(paragraphs["paragraph_extra_gap_median_pt"] - 5.5) <= 2.0
    title_available = bool(title and title_style.get("font_size_pt") is not None and title_style.get("bold_ratio") is not None and title_box and front["horizontal_rules"])
    title_ok = title_available and abs(title_style["font_size_pt"] - 17) <= 0.5 and title_style["bold_ratio"] >= 0.9 and (front.get("title_alignment_delta_pt") or 999) <= 8 and title_rules_ok
    abstract_available = bool(abstract and abstract.get("body") and abstract.get("left_indent_pt") is not None and abstract.get("right_indent_pt") is not None and abstract["style"].get("font_size_pt") is not None and abstract["style"].get("baseline_gap_pt") is not None and abstract["heading_style"].get("font_size_pt") is not None and abstract["heading_style"].get("bold_ratio") is not None and abstract.get("heading_alignment_delta_pt") is not None)
    abstract_ok = abstract_available and abs(abstract["left_indent_pt"] - 36) <= 5 and abs(abstract["right_indent_pt"] - 36) <= 5 and abs(abstract["style"]["font_size_pt"] - 10) <= 0.5 and abs(abstract["style"]["baseline_gap_pt"] - 11) <= 0.75 and abstract["paragraph_count"] == 1 and abs(abstract["heading_style"]["font_size_pt"] - 12) <= 0.5 and abstract["heading_style"]["bold_ratio"] >= 0.9 and abstract["heading_alignment_delta_pt"] <= 8
    h1_available = bool(h1 and all(item["style"].get("font_size_pt") is not None and item.get("bbox") for item in h1))
    h1_ok = h1_available and all(abs(item["style"]["font_size_pt"] - 12) <= 0.5 and (item["style"].get("bold_ratio") or 0) >= 0.8 and abs(item["bbox"][0] - geometry["left_pt"]) <= 3 for item in h1)
    rule_results = [
        result("NIPS-01", "正文区域几何", measured_status(geometry_available, geometry_ok), {"body_geometry": geometry}),
        result("NIPS-02", "正文字体与行距", measured_status(body_style_available, body_font_ok and baseline_ok and font_ok), {"body_font_size_pt": geometry["font_size_mode_pt"], "font_name": geometry["font_name_mode"], "times_compatible": font_ok, "baseline_gap_pt": geometry["baseline_gap_median_pt"]}),
        result("NIPS-03", "段落格式", measured_status(paragraph_available, paragraph_ok), {"paragraph_metrics": paragraphs}),
        result("NIPS-04", "论文标题", measured_status(title_available, title_ok), {"title": front}),
        result("NIPS-05", "摘要", measured_status(abstract_available, abstract_ok), {"abstract": abstract}),
        result("NIPS-06", "一级标题", measured_status(h1_available, h1_ok), {"headings": h1}),
        result("NIPS-07", "二级标题", "not_applicable" if not h2 else measured_status(all(item["style"].get("font_size_pt") is not None for item in h2), all(item["style"].get("font_size_pt") is not None and abs(item["style"]["font_size_pt"] - 10) <= 0.5 for item in h2)), {"headings": h2}),
        result("NIPS-08", "三级标题", "not_applicable" if not h3 else measured_status(all(item["style"].get("font_size_pt") is not None for item in h3), all(item["style"].get("font_size_pt") is not None and abs(item["style"]["font_size_pt"] - 10) <= 0.5 for item in h3)), {"headings": h3}),
        result("NIPS-09", "参考文献标题", measured_status(bool(ref and any(item["role"] == "reference_entry" for item in blocks)), bool(ref and text_key(ref["text"]) in REFERENCE_NAMES and not NUMBERED_HEADING_RE.match(ref["text"]) and all((item["page_number"], item["bbox"][1]) > (ref["page_number"], ref["bbox"][1]) for item in blocks if item["role"] == "reference_entry" and item.get("bbox")))), {"references_heading": ref, "reference_entry_count": sum(1 for item in blocks if item["role"] == "reference_entry" and ref and item["page_number"] >= ref["page_number"])}),
    ]
    figure = captions["figures"]
    table = captions["tables"]
    kinds = [item for item in (figure, table) if item["count"]]
    caption_ok = all(
        item["continuous"]
        and all(
            record["position"] == EXPECTED_CAPTION_POSITION[kind]
            and record["centered"]
            for record in item["items"]
        )
        for kind, item in (("figures", figure), ("tables", table))
        if item["count"]
    )
    rule_results.append(result("NIPS-10", "图表结构", "compliant" if kinds and caption_ok else "non_compliant" if kinds else "not_applicable", {"figures": figure, "tables": table}))
    return {"body_geometry": geometry, "front_matter": front, "abstract": abstract, "paragraph_metrics": paragraphs, "headings": headings, "captions": captions, "references": {"heading": ref}, "rules": rule_results}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--pdf", type=Path)
    parser.add_argument("--mineru-json", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--rule-document", type=Path)
    parser.add_argument("--self-test", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.self_test:
        run_self_tests()
        print("NeurIPS rule-evidence self-tests passed.")
        if not args.pdf and not args.mineru_json and not args.output:
            return 0
    if not args.pdf or not args.pdf.is_file():
        raise SystemExit(f"PDF not found: {args.pdf}")
    if not args.mineru_json or not args.mineru_json.is_file():
        raise SystemExit(f"MinerU JSON not found: {args.mineru_json}")
    if not args.output:
        raise SystemExit("--output is required when building evidence")
    native = extract_pymupdf(args.pdf)
    payload = json.loads(args.mineru_json.read_text(encoding="utf-8"))
    semantic = parse_mineru(payload, native["pages"])
    fused = fuse_blocks(native, semantic)
    report = evaluate(native, fused)
    derived_facts = build_derived_facts(native, report, fused)
    fused_blocks = serialize_fused_blocks(fused)
    output = {
        "schema_version": SCHEMA_VERSION,
        "source": {
            "pdf": str(args.pdf.resolve()),
            "pdf_sha256": sha256(args.pdf),
            "mineru": str(args.mineru_json.resolve()),
            "precedence": {
                "geometry_and_style": "pymupdf",
                "semantic_roles": "mineru",
                "semantic_text": "mineru",
                "object_container_geometry": "mineru_cross_checked_by_pymupdf",
            },
        },
        "quality": {
            "mineru_block_count": len(fused),
            "matched_block_count": sum(item["match_status"] == "matched" for item in fused),
            "unmatched_block_count": sum(item["match_status"] != "matched" for item in fused),
            "pymupdf_span_count": len(native["spans"]),
            "pymupdf_line_count": len(native["lines"]),
            "pymupdf_object_count": len(native["objects"]),
        },
        "derived_facts": derived_facts,
        "rule_evidence_index": build_rule_evidence_index(report["rules"]),
        "fused_blocks": fused_blocks,
        "native_facts": {"pages": native["pages"], "spans": native["spans"], "lines": native["lines"], "objects": native["objects"], "drawings": native["drawings"]},
    }
    output = round_floats(output)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.rule_document:
        write_rule_document(args.rule_document, output)
    print(
        json.dumps(
            {
                "output": str(args.output.resolve()),
                "rule_document": str(args.rule_document.resolve()) if args.rule_document else None,
                "quality": output["quality"],
                "rules": [
                    {"rule_id": item["rule_id"], "title": item["title"], "result": item["result"]}
                    for item in output["derived_facts"]["rule_assessments"]
                ],
            },
            ensure_ascii=True,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
